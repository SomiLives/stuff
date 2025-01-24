from flask import Flask, request, jsonify
from langchain_community.document_loaders import PyPDFLoader, DirectoryLoader, TextLoader, CSVLoader
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain.prompts import PromptTemplate
from langchain.callbacks.manager import CallbackManager
from langchain.chains import RetrievalQA
from langchain.callbacks import StreamingStdOutCallbackHandler
from langchain_community.llms import Ollama
from docx import Document as DocxDocument
import os
import shutil
from langchain.schema import Document

# Initialize Flask app
app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # Max file size (50MB)

# Paths
DATA_PATH = 'data/'
DB_FAISS_PATH = 'vectorstore/'

# Ensure directories exist
os.makedirs(DATA_PATH, exist_ok=True)
os.makedirs(DB_FAISS_PATH, exist_ok=True)

# Custom prompt template
custom_prompt_template = """
You are tasked with answering questions strictly based on the provided context and you provide detailed explanation.Provide the whole details on the question you are asked. Be thorough and do not omit any important details.
- Do not use external knowledge or make assumptions outside the context.
- If the context does not contain enough information, respond with "The context does not provide sufficient information to answer this question."
Context: {context}
Question: {question}
Answer (clear, concise, and detailed):
"""

# Custom loader for .docx files
class DocxLoader:
    def __init__(self, path):
        self.path = path

    def load(self):
        """Loads the DOCX file and returns the content as a list of documents."""
        try:
            doc = DocxDocument(self.path)
            content = "\n".join([para.text for para in doc.paragraphs])
            # Return as LangChain's Document type with the page_content attribute
            return [Document(page_content=content)]
        except Exception as e:
            print(f"Error loading DOCX file {self.path}: {e}")
            return []

    def lazy_load(self):
        """Returns a generator to lazily load the document."""
        yield from self.load()

def determine_chunk_size(doc_length, max_size=1200, min_size=500):
    """Determine optimal chunk size based on document length."""
    if doc_length > 5000:
        return max_size
    elif doc_length < 1000:
        return min_size
    else:
        return doc_length

def create_vector_db():
    """Creates and saves the FAISS vector database."""
    print("Creating vector database...")

    # Loaders for different file types, including the new .docx loader
    loaders = [
        DirectoryLoader(DATA_PATH, glob='*.pdf', loader_cls=PyPDFLoader),
        DirectoryLoader(DATA_PATH, glob='*.txt', loader_cls=TextLoader),
        DirectoryLoader(DATA_PATH, glob='*.csv', loader_cls=CSVLoader),
        DirectoryLoader(DATA_PATH, glob='*.docx', loader_cls=DocxLoader),  # New loader for DOCX files
    ]

    documents = []
    for loader in loaders:
        try:
            docs = loader.load()
            print(f"Loaded {len(docs)} documents using {loader.__class__.__name__}")
            for doc in docs:
                doc_size = len(doc.page_content)
                print(f"Document size: {doc_size} characters")
            documents.extend(docs)
        except Exception as e:
            print(f"Error loading documents with {loader.__class__.__name__}: {e}")

    if not documents:
        print("No documents found to process.")
        return None

    # Adjusted to use `documents` size for chunk size calculation
    avg_doc_size = sum(len(doc.page_content) for doc in documents) / len(documents)
    print(f"Average document size: {avg_doc_size:.2f} characters")

    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=int(determine_chunk_size(len(documents))),
        chunk_overlap=10
    )
    texts = text_splitter.split_documents(documents)
    print(f"Total text chunks: {len(texts)}")

    try:
        embeddings = HuggingFaceEmbeddings(
            model_name='sentence-transformers/all-MiniLM-L6-v2',
            model_kwargs={'device': 'cpu'}
        )
        db = FAISS.from_documents(texts, embeddings)
        db.save_local(DB_FAISS_PATH)
        print("Vector database created and saved successfully!")
        return db
    except Exception as e:
        print(f"Error creating vector database: {e}")
        return None



def set_custom_prompt():
    """Sets the Q and A retrieval prompt."""
    try:
        prompt = PromptTemplate(template=custom_prompt_template, input_variables=['context', 'question'])
        return prompt
    except Exception as e:
        print(f"Error initializing custom prompt: {e}")
        return None


def load_llm():
    """Loads the LLM (OLLAMA) model."""
    try:
        llm = Ollama(
            model="llama3.2",
            temperature=0,
            verbose=True,
            callback_manager=CallbackManager([StreamingStdOutCallbackHandler()])
        )
        return llm
    except Exception as e:
        print(f"Error loading LLM: {e}")
        return None


def load_db():
    """Loads the FAISS vector database."""
    embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2",
                                       model_kwargs={'device': 'cpu'})
    try:
        db = FAISS.load_local(DB_FAISS_PATH, embeddings, allow_dangerous_deserialization=True)
        return db
    except Exception as e:
        print(f"Error loading FAISS index: {e}")
        return None


def retrieval_qa_chain(llm, prompt, db):
    """Sets up the retrieval QA chain."""
    try:
        retriever = db.as_retriever(search_type="similarity", search_kwargs={'k': 2})
        qa_chain = RetrievalQA.from_chain_type(
            llm=llm,
            chain_type='stuff',
            retriever=retriever,
            return_source_documents=True,
            chain_type_kwargs={'prompt': prompt}
        )
        return qa_chain
    except Exception as e:
        print(f"Error setting up Retrieval QA chain: {e}")
        return None


def qa_bot():
    """Sets up the question-answering bot with retrieval."""
    llm = load_llm()
    db = load_db()

    if not llm or not db:
        return None

    qa_prompt = set_custom_prompt()
    if not qa_prompt:
        return None

    return retrieval_qa_chain(llm, qa_prompt, db)


@app.route('/clear_data_and_vector-db', methods=['DELETE'])
def delete_contents():
    """API endpoint to delete all contents in the vectorstore and data directories."""
    try:
        # Delete contents of the vectorstore directory
        if os.path.exists(DB_FAISS_PATH):
            for file in os.listdir(DB_FAISS_PATH):
                file_path = os.path.join(DB_FAISS_PATH, file)
                if os.path.isfile(file_path) or os.path.islink(file_path):
                    os.unlink(file_path)  # Remove file or symlink
                elif os.path.isdir(file_path):
                    shutil.rmtree(file_path)  # Remove directory
            print("Vectorstore directory cleared.")

        # Delete contents of the data directory
        if os.path.exists(DATA_PATH):
            for file in os.listdir(DATA_PATH):
                file_path = os.path.join(DATA_PATH, file)
                if os.path.isfile(file_path) or os.path.islink(file_path):
                    os.unlink(file_path)  # Remove file or symlink
                elif os.path.isdir(file_path):
                    shutil.rmtree(file_path)  # Remove directory
            print("Data directory cleared.")

        return jsonify({"message": "Contents of vectorstore and data directories deleted successfully."}), 200
    except Exception as e:
        print(f"Error deleting directory contents: {e}")
        return jsonify({"error": f"Failed to delete contents: {str(e)}"}), 500


@app.route('/upload_new_file', methods=['POST'])
def upload_file():
    """API endpoint to upload a file."""
    if 'file' not in request.files:
        return jsonify({"error": "No file provided"}), 400

    file = request.files['file']
    file_path = os.path.join(DATA_PATH, file.filename)
    try:
        file.save(file_path)
        return jsonify({"message": f"File {file.filename} uploaded successfully!"}), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/insert_updated_file', methods=['PUT'])
def update():
    """API endpoint to update an uploaded file."""
    try:
        # List all file names in the data directory
        try:
            file_names = [file for file in os.listdir(DATA_PATH) if os.path.isfile(os.path.join(DATA_PATH, file))]
        except Exception as e:
            print(f"Error listing files in directory '{DATA_PATH}': {e}")
            return jsonify({"error": f"Unable to access data directory. {str(e)}"}), 500

        # Retrieve the uploaded file
        file = request.files['file']
        uploaded_filename = file.filename
        file_path = os.path.join(DATA_PATH, uploaded_filename)

        # Check if the uploaded file exists in the data directory
        if uploaded_filename not in file_names:
            return jsonify({"error": f"File '{uploaded_filename}' does not exist in the data directory. Cannot update."}), 404

        # If file exists, overwrite it with the uploaded file
        file.save(file_path)
        print(f"File '{uploaded_filename}' updated successfully.")

        # Recreate the vector database after the file is updated
        create_vector_db()

        return jsonify({"message": f"File '{uploaded_filename}' updated and processed successfully!"}), 200

    except Exception as e:
        print(f"Error during file update: {e}")
        return jsonify({"error": f"An error occurred: {str(e)}"}), 500


@app.route('/user_query_to_get_response', methods=['POST'])
def api_get_answer():
    """API endpoint to get an answer from the QA system."""
    data = request.json
    user_query = data.get("user_query")
    qa_result = qa_bot()
    if not qa_result:
        return jsonify({"error": "QA system initialization failed."}), 500
    response_from_model = qa_result({'user_query': user_query})
    return jsonify({"response_from_model": response_from_model['response_from_model']}), 200






